# -*- coding: utf-8 -*-
"""更新 南丰蜜桔模型管理系统-需求说明书.docx：
1) 第一章 项目概述 重写为参考整改实施方案；
2) 全部截图替换为新生成的正确分辨率 PNG（后台/大屏 1920x1080，移动端 430x932）；
3) 补充 HTML 中已存在但文档缺失的 7 个页面；
4) 更新附录页面清单与计数。
"""
import os, copy, glob
from docx import Document
from docx.shared import Cm, RGBColor
from docx.oxml.ns import qn

SRC = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统-需求说明书.docx"
OUT = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统-需求说明书.docx"
SHOT = r"D:\dev\pro_orangemodel\docs\screenshots"

doc = Document(SRC)
body = doc.element.body

# ---------- 截图映射 ----------
shot_files = glob.glob(os.path.join(SHOT, "*.png"))
shot_names = {os.path.basename(f) for f in shot_files}
def shot_for(page):
    """page 形如 backhand/growth/growth_model.html -> 文件名"""
    name = page.replace("/", "__").rsplit(".", 1)[0] + ".png"
    if name in shot_names:
        return os.path.join(SHOT, name)
    # 失效引用回退
    if page.endswith("growth_model_memo.html"):
        return shot_for("backhand/growth/growth_model.html")
    if page.endswith("pest_model_memo.html"):
        return shot_for("backhand/pest/pest_model.html")
    return None

# ---------- 样式探测 ----------
def para_by_text(startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    return None

def style_of(startswith):
    p = para_by_text(startswith)
    return p.style.name if p else "Normal"

sub3_style = style_of("3.1.1")          # 三级标题样式
ch1sub_style = style_of("1.2")          # 第一章子标题样式
body_style = None
p10 = para_by_text("1.1")
if p10 is not None:
    paras = doc.paragraphs
    idx = next(i for i, pp in enumerate(paras) if pp._p is p10._p)
    for q in paras[idx+1:idx+4]:
        if q.text.strip() and q.style.name != ch1sub_style:
            body_style = q.style.name; break
if body_style is None:
    body_style = "Normal"

print("styles: sub3=", sub3_style, "ch1sub=", ch1sub_style, "body=", body_style)

# ---------- 工具 ----------
def new_para(text="", style=None, size=None, bold=False):
    p = doc.add_paragraph()
    if style:
        try: p.style = style
        except Exception: pass
    if text:
        r = p.add_run(text)
        if size: r.font.size = size
        r.font.bold = bold
    return p

def new_table(rows, cols, header=None):
    t = doc.add_table(rows=rows, cols=cols)
    try: t.style = "Table Grid"
    except Exception: pass
    if header:
        for j, h in enumerate(header):
            c = t.rows[0].cells[j]
            c.text = ""
            rr = c.paragraphs[0].add_run(h)
            rr.font.bold = True
    return t

# ===================================================================
# 1) 重写第一章 项目概述
# ===================================================================
ch1 = para_by_text("第一章 项目概述")
ch2 = para_by_text("第二章 业务流程")
assert ch1 is not None and ch2 is not None, "章节标题未找到"

children = list(body)
i1 = children.index(ch1._p)
i2 = children.index(ch2._p)
for c in children[i1+1:i2]:
    body.remove(c)

ch1_blocks = []
def H(text):
    p = new_para(text, style=ch1sub_style); ch1_blocks.append(p._p)
def B(text):
    p = new_para(text, style=body_style); ch1_blocks.append(p._p)

H("1.1 项目背景")
B("南丰蜜桔是南丰县特色优势农业产业，是助力地方乡村振兴、农业增效、农户增收的核心支柱产业。当前南丰蜜桔种植产业仍存在传统种植模式粗放、生长环境管控不精准、农事管理缺乏科学依据、病虫害防控滞后等问题。为破解产业发展瓶颈，依据《关于南丰县蜜桔全产业链智慧建设项目技术审查的情况说明》提出的整改意见，本项目在现有南丰蜜桔全产业链大数据平台的数据基底和业务应用基础上，依托现有果园物联网监测设备、气象监测系统等基础硬件资源，结合人工智能、物联网、大数据等现代信息技术，对南丰蜜桔模型管理系统进行完善与整改，推动果园管理由依靠经验向数据辅助转变。")
H("1.2 整改与建设原则")
B("本次整改坚持充分利用现有平台、设备和数据，不新建独立平台、不大规模新增物联网设备，重点完善生长模型、病虫害模型和销售数据收集三项功能。所有整改内容均整合到现有模型管理系统后台、模型监测大屏、蜜桔农业助手和第三方对接能力中，对现有功能进行补充、整合和打通，形成统一的应用体系。")
H("1.3 建设基础（现有四大能力）")
B("本系统以现有建设成果为建设基础，包含四大能力：（1）模型管理后台——面向农业农村局业务管理人员，提供模型参数配置、运算调度、规则维护、预警处置与系统管理；（2）模型监测大屏——面向产业研判与成果展示，可视化呈现生长模型与病虫害模型的运行态势；（3）蜜桔农业助手——面向种植端用户的微信小程序，提供环境监测、数据上报、AI病虫害识别、长势评价与农事建议；（4）第三方对接——保留用户单点登录、设备（IoT）数据接入及外部数据（气象局/农业大数据/价格行情等）抓取能力。")
H("1.4 建设内容与范围")
B("本次完善围绕三项核心功能展开：（1）完善生长模型——补充模型说明、基础档案、因子采集、模型计算、诊断分析、农事建议和评价报告等功能，采用适宜度、阈值门控、扣分与综合评分方法，输出综合评分、生长状态等级、主要扣分项、问题原因、农事建议及建议处理时间；（2）完善病虫害模型——结合天气、温湿度、生长阶段与历史发生记录，对黄龙病、炭疽病、黑点病、疮痂病、红蜘蛛、柑桔木虱等主要病虫害进行低/中/高三级风险提示，并提供叶片、枝条、果实、虫体图片的AI辅助识别；（3）完善销售数据收集——建设销售主体填报与公开网站采集相结合的价格、市场行情与产销信息采集统计功能。详细功能范围见第三章。")
H("1.5 实施周期与试点")
B("本项目选择少量具有代表性的果园进行试运行，邀请农技人员对评分结果、病虫害预警和农事建议进行确认，并根据试运行情况对指标范围和建议规则进行调整。项目建设周期控制在2—3个月，完成系统测试、人员培训和验收资料编制后上线使用。")

# 插入到 ch1 标题之后
cur = ch1._p
for el in ch1_blocks:
    cur.addnext(el); cur = el

# ===================================================================
# 2) 替换全部截图
# ===================================================================
def para_maps_to_page(p_el):
    # 向上回溯找 对应文件：
    paras = list(body.iter(qn('w:p')))
    idx = paras.index(p_el)
    for k in range(idx-1, max(-1, idx-9), -1):
        txt = "".join(n.text or "" for n in paras[k].iter(qn('w:t')))
        if "对应文件：" in txt:
            return txt.split("对应文件：", 1)[1].strip().split()[0]
    return None

imgs_replaced = 0
imgs_missing = 0
current_page = None
for el in list(body):
    if el.tag == qn('w:p'):
        txt = "".join(n.text or "" for n in el.iter(qn('w:t')))
        if "对应文件：" in txt:
            current_page = txt.split("对应文件：", 1)[1].strip().split()[0]
        if el.find('.//'+qn('w:drawing')) is not None and current_page:
            png = shot_for(current_page)
            if not png:
                imgs_missing += 1
            else:
                width_cm = 6.8 if current_page.startswith("mobile/") else 15.0
                tmp = doc.add_paragraph()
                run = tmp.add_run()
                run.add_picture(png, width=Cm(width_cm))
                new_drawing = tmp._p.find('.//'+qn('w:drawing'))
                old = el.find('.//'+qn('w:drawing'))
                old.getparent().replace(old, copy.deepcopy(new_drawing))
                tmp._p.getparent().remove(tmp._p)
                imgs_replaced += 1

print("imgs replaced=", imgs_replaced, "missing=", imgs_missing)

# 清理未被引用的旧图片关系（去除错误分辨率的残留图）
referenced = set()
for drawing in body.iter(qn('w:drawing')):
    blip = drawing.find('.//'+qn('a:blip'))
    if blip is not None:
        referenced.add(blip.get(qn('r:embed')))
for rid in list(doc.part.rels):
    rel = doc.part.rels[rid]
    if rid not in referenced and rel.reltype.endswith("/image"):
        doc.part.drop_rel(rid)
print("orphan image rels dropped")

# ===================================================================
# 3) 补充缺失页面
# ===================================================================
MISSING = [
    {"anchor":"backhand/alert/send_template.html","num":"3.1.6.5 内容设置",
     "file":"backhand/alert/content_setting.html",
     "desc":"预警模板的内容设置子功能，配置短信、APP与微信三类渠道的发送内容，采用富文本编辑器编辑模板正文。",
     "fields":[("模板名称","文本","模板标识名称"),("适用级别","下拉","红色/橙色/黄色/蓝色/普通"),
               ("短信内容","文本","短信渠道发送正文"),("APP内容","富文本","APP推送正文（Quill编辑）"),
               ("微信内容","富文本","微信模板消息正文"),("变量占位符","文本","支持插入动态变量")]},
    {"anchor":"backhand/alert/send_template.html","num":"3.1.6.6 流程设置",
     "file":"backhand/alert/flow_setting.html",
     "desc":"预警模板的流程设置子功能，配置审批流程、发送流程与失败重试机制。",
     "fields":[("审批流程","文本","预警发布前审批节点"),("发送流程","文本","渠道发送顺序"),
               ("重试机制","开关","发送失败是否自动重试"),("重试次数","数值","最大重试次数"),
               ("重试间隔","数值","重试时间间隔（秒）")]},
    {"anchor":"backhand/system/role_manage.html","num":"3.1.8.3 模型配置",
     "file":"backhand/system/model_config.html",
     "desc":"系统级模型运行配置，管理生长模型与病虫害模型的运算开关、计算周期与参数版本。",
     "fields":[("生长模型开关","开关","生长模型计算启停"),("病虫害模型开关","开关","病虫害模型计算启停"),
               ("计算周期","下拉","每日/每周/手动"),("数据校准开关","开关","是否启用参数校准"),
               ("参数版本","文本","当前模型参数版本号")]},
    {"anchor":"backhand/wechat/recognize_log.html","num":"3.1.7.6 人工上报（微信后台）",
     "file":"backhand/wechat/manual_collect.html",
     "desc":"微信后台侧的人工采集数据管理，审核农户/农技人员通过小程序上报的土壤、气象、农事与病虫数据。",
     "fields":[("数据ID","自动生成","如DC-001"),("采集人","文本","上报用户"),
               ("地块","下拉","关联果园地块"),("采集类型","下拉","土壤养分/土壤墒情/温湿度/农事操作"),
               ("采集时间","时间","数据时间"),("数据详情","文本","按类型动态展示"),("状态","标签","待审核/已通过/已驳回")]},
    {"anchor":"backhand/wechat/recognize_log.html","num":"3.1.7.7 问题管理",
     "file":"backhand/wechat/question_manage.html",
     "desc":"管理小程序用户提交的问题反馈与建议，跟踪处理状态与回复内容。",
     "fields":[("问题ID","自动生成","问题编号"),("用户","文本","提交用户"),
               ("问题类型","下拉","问题反馈/功能建议/使用疑问/其他"),("问题描述","文本","问题内容"),
               ("提交时间","时间","提交时间"),("处理状态","标签","待处理/处理中/已回复"),("处理回复","文本","官方回复内容")]},
    {"anchor":"mobile/mb_disease_ai.html","num":"3.2.5.1 识别结果页",
     "file":"mobile/mb_disease_result.html",
     "desc":"AI病虫害识别结果展示页，返回识别名称、置信度、风险等级与防治建议清单。",
     "fields":[("识别图片","图片","上传的原图"),("病名","文本","识别结果名称"),
               ("置信度","文本","识别置信度百分比"),("风险等级","标签","低/中/高风险"),
               ("防治建议","列表","4条推荐防治措施"),("重新识别","按钮","返回重新上传")]},
    {"anchor":"mobile/mb_disease_knowledge.html","num":"3.2.6.1 病害/虫害详情页",
     "file":"mobile/mb_disease_detail.html",
     "desc":"病虫害知识详情页，展示单一病虫害的描述、症状、防治方法、推荐用药与注意事项。",
     "fields":[("名称","文本","病虫害名称"),("类型","标签","病害/虫害"),
               ("症状","文本","症状表现描述"),("防治方法","文本","防治措施"),
               ("推荐用药","标签列表","drug-tag列表"),("注意事项","文本","使用注意")]},
]

# 按 anchor 分组
groups = {}
for m in MISSING:
    groups.setdefault(m["anchor"], []).append(m)

def find_image_para_for_page(page):
    for p_el in list(body.iter(qn('w:p'))):
        if p_el.find('.//'+qn('w:drawing')) is None:
            continue
        if para_maps_to_page(p_el) == page:
            return p_el
    return None

for anchor_page, blocks in groups.items():
    anchor_p = find_image_para_for_page(anchor_page)
    if anchor_p is None:
        print("WARN anchor not found:", anchor_page); continue
    cur = anchor_p
    for b in blocks:
        png = shot_for(b["file"])
        els = []
        h = new_para(b["num"], style=sub3_style); els.append(h._p)
        fp = new_para("对应文件：" + b["file"]); els.append(fp._p)
        cap = new_para("图：" + b["num"] + " - 主界面"); els.append(cap._p)
        if png:
            ip = doc.add_paragraph()
            ir = ip.add_run(); ir.add_picture(png, width=Cm(6.8 if b["file"].startswith("mobile/") else 15.0))
            els.append(ip._p)
        dp = new_para(b["desc"], style=body_style); els.append(dp._p)
        if b["fields"]:
            t = new_table(len(b["fields"])+1, 3, header=["字段名称","类型","说明"])
            for i, (n, ty, de) in enumerate(b["fields"], start=1):
                t.rows[i].cells[0].text = n
                t.rows[i].cells[1].text = ty
                t.rows[i].cells[2].text = de
            els.append(t._tbl)
        for el in els:
            cur.addnext(el); cur = el

print("missing pages inserted")

# ===================================================================
# 4) 更新附录页面清单与计数
# ===================================================================
appendix_adds = {
    "backhand": [("backhand/alert/content_setting.html","预警模板-内容设置"),
                 ("backhand/alert/flow_setting.html","预警模板-流程设置"),
                 ("backhand/system/model_config.html","系统-模型配置"),
                 ("backhand/wechat/manual_collect.html","微信后台-人工上报"),
                 ("backhand/wechat/question_manage.html","微信后台-问题管理")],
    "mobile": [("mobile/mb_disease_result.html","AI识别结果页"),
               ("mobile/mb_disease_detail.html","病害/虫害详情页")],
}
# 找到附录表（table 83 backhand, 84 mobile）并按匹配追加行
for ti, t in enumerate(doc.tables):
    first = t.rows[0].cells[0].text.strip() if t.rows else ""
    # 通过表头判断：页面路径 / 功能描述
    hdr = ""
    if t.rows:
        hdr = t.rows[0].cells[0].text.strip() + "/" + (t.rows[0].cells[1].text.strip() if t.columns.__len__()>1 else "")
    if hdr == "页面路径/功能描述":
        # 判断属于哪端：看已有行是否含 backhand 或 mobile
        txt_all = "\n".join(c.text for r in t.rows for c in r.cells)
        if "backhand/" in txt_all and "mobile/" not in txt_all:
            key = "backhand"
        elif "mobile/" in txt_all:
            key = "mobile"
        else:
            key = None
        if key and key in appendix_adds:
            for path, desc in appendix_adds[key]:
                row = t.add_row()
                row.cells[0].text = path
                row.cells[1].text = desc

# 更新计数标签
for p in doc.paragraphs:
    if "模型后台管理系统（" in p.text:
        p.text = p.text.replace(p.text[p.text.index("（"):], "（46个页面）")
    if "微信小程序（" in p.text:
        p.text = p.text.replace(p.text[p.text.index("（"):], "（22个页面）")

# ===================================================================
# 5) 全文档字体置黑
# ===================================================================
for p in doc.paragraphs:
    for r in p.runs:
        r.font.color.rgb = RGBColor(0,0,0)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0,0,0)

doc.save(OUT)
print("SAVED:", OUT)
