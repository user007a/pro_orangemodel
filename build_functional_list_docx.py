# -*- coding: utf-8 -*-
"""将 南丰蜜桔模型管理系统项目-功能清单.xlsx 转成 Word 版（表格形式，全黑配色）。"""
import os, zipfile, re, xml.etree.ElementTree as ET
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统项目-功能清单.xlsx"
OUT = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统-功能清单.docx"

# ---------- 1. 解析 xlsx ----------
ns = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
z = zipfile.ZipFile(SRC)
ss = {}
root = ET.fromstring(z.read("xl/sharedStrings.xml"))
for i, t in enumerate(root.findall(f"{ns}si")):
    ss[i] = "".join(node.text or "" for node in t.iter(f"{ns}t"))

def cr(ref):
    m = re.match(r"([A-Z]+)(\d+)", ref); return m.group(1), int(m.group(2))
def ci(c):
    n = 0
    for ch in c: n = n * 26 + (ord(ch) - 64)
    return n
def cv(c):
    t = c.get("t"); v = c.find(f"{ns}v"); isn = c.find(f"{ns}is")
    if t == "s" and v is not None: return ss.get(int(v.text), "")
    if t == "inlineStr" and isn is not None: return "".join(x.text or "" for x in isn.iter(f"{ns}t"))
    if v is not None: return v.text or ""
    return ""

root = ET.fromstring(z.read("xl/worksheets/sheet1.xml"))
rows = {}
for c in root.iter(f"{ns}c"):
    col, r = cr(c.get("r")); rows.setdefault(r, {})[ci(col)] = cv(c)

records = []  # [系统名称, 一级菜单, 二级菜单, 内容描述]
for r in range(3, 82):  # 跳过标题行(1)与表头行(2)
    d = rows.get(r, {})
    sys_v = d.get(1, "").strip()
    l1 = d.get(2, "").strip()
    l2 = d.get(3, "").strip()
    desc = d.get(4, "").strip()
    if not l2 and not sys_v and not l1:
        continue
    records.append([sys_v, l1, l2, desc])

# ---------- 2. 写 Word ----------
def set_run_font(run, size=10.5, bold=False, name="宋体", eastasia="宋体"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), eastasia)

def fill(cell, text, size=10.5, bold=False, name="宋体", eastasia="宋体", align=None):
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, name=name, eastasia=eastasia)
    # 单倍行距，避免表格过高
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing"); pPr.append(spacing)
    spacing.set(qn("w:line"), "276"); spacing.set(qn("w:lineRule"), "auto")

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.0)

# 标题
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(title.add_run("南丰蜜桔模型管理系统 — 功能清单"),
             size=16, bold=True, name="黑体", eastasia="黑体")
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(sub.add_run("（软件功能范围说明，与《软件设计方案》《软件技术方案》配套使用）"),
             size=10, name="宋体", eastasia="宋体")

# 表格
table = doc.add_table(rows=1, cols=4)
try:
    table.style = "Table Grid"
except Exception:
    pass
headers = ["系统名称", "一级菜单", "二级菜单", "内容描述"]
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    fill(hdr_cells[i], h, size=10.5, bold=True, name="黑体", eastasia="黑体",
         align=WD_ALIGN_PARAGRAPH.CENTER)

widths = [Cm(2.6), Cm(2.8), Cm(3.4), Cm(8.0)]

for rec in records:
    cells = table.add_row().cells
    fill(cells[0], rec[0])
    fill(cells[1], rec[1])
    fill(cells[2], rec[2])
    fill(cells[3], rec[3])

# 设置列宽
for ci_, w in enumerate(widths):
    for cell in table.columns[ci_].cells:
        cell.width = w

# 纵向合并：子系统 / 一级菜单
n = len(records)
def merge_col(col):
    i = 0
    while i < n:
        j = i
        while j + 1 < n and records[j + 1][col] == records[i][col] \
              and (col == 0 or records[j + 1][0] == records[i][0]):
            j += 1
        if j > i:
            table.cell(i + 1, col).merge(table.cell(j + 1, col))
        i = j + 1
merge_col(0)
merge_col(1)

# 合并单元格内文字垂直居中
for row in table.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        vAlign = tcPr.find(qn("w:vAlign"))
        if vAlign is None:
            vAlign = OxmlElement("w:vAlign"); tcPr.append(vAlign)
        vAlign.set(qn("w:val"), "center")

doc.save(OUT)
print("saved:", OUT, "| records:", n)
