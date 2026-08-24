# -*- coding: utf-8 -*-
"""增强需求说明书：第二章加流程图；各多Tab页展开Tab截图与描述；生长大屏补充KPI弹窗截图与描述。"""
import json, os
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOCX = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统-需求说明书.docx"
SHOT = r"D:\dev\pro_orangemodel\docs\screenshots"
FLOW = r"D:\dev\pro_orangemodel\docs\flowcharts"

doc = Document(DOCX)

# ---------- 字体/样式工具 ----------
def set_run(run, size=10.5, bold=False, name="宋体", eastasia="宋体", color="000000"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run.font.color.rgb = __import__("docx").shared.RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts"); rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), eastasia)

def insert_para_after(ref, text=None, size=10.5, bold=False, align=None, color="000000"):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text is not None:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    ref._p.addnext(p._p)
    return p

def add_image_after(ref, path, width_cm):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    ref._p.addnext(p._p)
    return p

def find_para(startswith):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    return None

def el_text(p_elem):
    return "".join(n.text or "" for n in p_elem.iter(qn("w:t"))).strip()

def find_docline(normpath):
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("对应文件："):
            pp = t.split("对应文件：", 1)[1].strip()
            if pp.replace("_memo", "") == normpath:
                return p
    return None

def next_nonempty_after(para):
    cur = para._p.getnext()
    while cur is not None:
        if cur.tag == qn("w:p"):
            if el_text(cur):
                return Paragraph(cur, doc)
        cur = cur.getnext()
    return None

def find_image_after(ref):
    cur = ref._p.getnext()
    while cur is not None:
        if cur.tag == qn("w:p"):
            if cur.find(".//" + qn("w:drawing")) is not None:
                return Paragraph(cur, doc)
        cur = cur.getnext()
    return ref

def shot_path(rel, suffix):
    base = rel.replace("/", "__")
    return os.path.join(SHOT, base + suffix)

# ---------- 1. 第二章业务流程配流程图 ----------
flow_map = [
    ("2.1", "flow_growth.png", "图：2.1 生长模型业务流程"),
    ("2.2", "flow_pest.png",   "图：2.2 病虫害模型业务流程"),
    ("2.3", "flow_alert.png",  "图：2.3 预警管理业务流程"),
    ("2.4", "flow_mobile.png", "图：2.4 移动端全链路业务流程"),
]
for hp, fname, cap in flow_map:
    h = find_para(hp)
    if not h:
        print("WARN heading not found:", hp); continue
    desc = next_nonempty_after(h)
    anchor = desc if desc else h
    ip = add_image_after(anchor, os.path.join(FLOW, fname), 15.0)
    insert_para_after(ip, cap, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    print("flow inserted:", hp)

# ---------- 2. 多Tab页展开 ----------
manifest = json.load(open(os.path.join(SHOT, "manifest.json"), encoding="utf-8"))
TAB_RELS = [k for k in manifest if not k.startswith("__")]
for rel in TAB_RELS:
    labels = manifest[rel]
    if len(labels) < 2:
        continue
    ref = find_docline(rel)
    if not ref:
        print("WARN docline not found for tab page:", rel); continue
    img = find_image_after(ref)
    last = img
    lead = "该页面包含 %d 个标签页：%s，默认展示“%s”，其余标签页内容如下。" % (
        len(labels), "、".join(labels), labels[0])
    last = insert_para_after(last, lead, size=10.5)
    for i in range(1, len(labels)):
        cap = "图：%s - %s" % (rel, labels[i])
        c = insert_para_after(last, cap, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        w = 6.8 if rel.startswith("mobile/") else 15.0
        im = add_image_after(c, shot_path(rel, "__tab%d.png" % (i + 1)), w)
        last = im
    print("tabs inserted:", rel, labels)

# ---------- 3. 生长大屏 KPI 弹窗 ----------
modal_caps = manifest.get("__growth_dashboard_modal", [])
if modal_caps:
    cap_ref = find_para("图：3.3.1 生长模型监测大屏 - 主界面")
    if not cap_ref:
        print("WARN dashboard caption not found"); 
    else:
        last = cap_ref
        last = insert_para_after(last,
            "大屏顶部为 8 张 KPI 统计卡片，点击任意卡片可弹出明细弹窗，展示趋势图表、统计指标与明细数据。各卡片弹窗内容如下：",
            size=10.5)
        for i, m in enumerate(modal_caps):
            cap = "图：3.3.1 生长模型监测大屏 - KPI弹窗（%s）" % m
            c = insert_para_after(last, cap, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            im = add_image_after(c, shot_path("dataanlye/growth_dashboard", "__modal%d.png" % (i + 1)), 15.0)
            last = im
        print("dashboard modals inserted:", len(modal_caps))

out = DOCX
doc.save(out)
print("SAVED", out)
