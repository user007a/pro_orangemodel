# -*- coding: utf-8 -*-
"""在已插入的 Tab 截图与 KPI 弹窗截图后追加一句功能说明（仅文本，不重复插图）。"""
import json, os, re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DOCX = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统-需求说明书.docx"
SHOT = r"D:\dev\pro_orangemodel\docs\screenshots"

doc = Document(DOCX)

def set_run(run, size=9.5, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "宋体"
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = __import__("docx").oxml.OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), "宋体"); rf.set(qn("w:hAnsi"), "宋体"); rf.set(qn("w:eastAsia"), "宋体")

def find_para_by_text(text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None

def el_text(e):
    return "".join(n.text or "" for n in e.iter(qn("w:t"))).strip()

def image_after(ref):
    cur = ref._p.getnext()
    while cur is not None:
        if cur.tag == qn("w:p") and cur.find(".//" + qn("w:drawing")) is not None:
            return Paragraph(cur, doc)
        cur = cur.getnext()
    return ref

def insert_after(ref, text, size=9.5):
    p = doc.add_paragraph()
    r = p.add_run(text); set_run(r, size=size)
    ref._p.addnext(p._p)
    return p

TAB_DESC = {
 "backhand/growth/growth_model.html": {
   "模型说明": "展示生长模型的总体说明、算法原理与核心参数配置，便于管理人员理解模型运行逻辑。"},
 "backhand/growth/base_archive.html": {
   "品系管理": "维护南丰蜜桔品种/品系基础信息，包括品种特性与适生条件。",
   "地块档案": "管理果园地块档案，记录地块位置、面积、株数等基础信息。",
   "阶段参数": "配置各生长阶段的模型参数与阶段划分规则。"},
 "backhand/growth/factor_collect.html": {
   "物联网实时数据": "展示由物联网设备实时回传的气象、墒情等监测数据。",
   "人工采样录入": "支持人工采样数据的录入与校核，补齐自动采集盲区。",
   "农事记录录入": "记录施肥、打药、修剪等农事操作数据，作为模型因子来源之一。"},
 "backhand/growth/farm_scheme_config.html": {
   "方案列表": "查看与维护农事建议方案清单。",
   "方案测试": "对方案进行模拟测试，验证建议的合理性与可行性。",
   "因子映射": "配置农事方案因子与模型因子的映射关系。"},
 "backhand/growth/calibration.html": {
   "历史数据对比": "将模型计算值与历史实测值进行对比分析，定位偏差。",
   "参数微调": "对模型参数进行人工微调，校准计算偏差。",
   "校准日志": "记录历次校准操作的时间、人员与调整内容。"},
 "backhand/pest/pest_data.html": {
   "气象数据 (Wt)": "展示气象监测数据，作为病虫害预测的基础因子。",
   "物候数据 (Pt)": "记录作物物候期数据，用于预测时间窗口修正。",
   "病情虫情 (Dt)": "汇总病害与虫情监测数据。",
   "虫源基数 (It)": "记录虫源基数监测数据，支撑发生量预测。",
   "历史数据 (Ht)": "提供历史病情虫情数据查询。"},
 "backhand/pest/pest_report.html": {
   "3天预测": "展示未来 3 天病虫害风险预测结果。",
   "7天预测": "展示未来 7 天病虫害风险预测结果。",
   "历史报告": "查看已生成的病虫害预测历史报告。"},
 "backhand/pest/pest_diagnosis.html": {
   "因子归因": "分析导致风险的主要影响因子及其贡献度。",
   "耦合分析": "展示多因子耦合作用对风险的影响。",
   "传播追踪": "追踪病害/虫害的空间传播路径与范围。"},
 "backhand/pest/pest_params.html": {
   "因子权重 (w)": "配置各预测因子的权重参数。",
   "耦合系数 (β/γ)": "配置因子间的耦合系数。",
   "时间窗口 (L/h)": "配置时间累积效应的窗口参数。",
   "θ权重 (θ1-θ5)": "配置各子模型的概率校准权重。"},
 "mobile/mb_data.html": {
   "环境指标": "移动端展示果园环境（气象、墒情）实时指标。",
   "虫情测报": "移动端展示虫情测报信息。"},
 "mobile/mb_data_input.html": {
   "土壤数据": "录入土壤检测数据。",
   "气象数据": "录入气象观测数据。",
   "农事记录": "记录农事操作。",
   "病虫上报": "上报病虫害发生情况。",
   "蜜桔销售": "录入蜜桔销售数据，支撑销售数据收集。"},
 "mobile/mb_messages.html": {
   "预警": "查看系统预警消息列表。",
   "风险": "查看病虫害风险消息。",
   "推荐": "查看农事推荐消息。",
   "农事": "查看农事提醒消息。"},
 "mobile/mb_profile_records.html": {
   "全部": "汇总展示个人上报的全部记录。",
   "土壤数据": "分类查看土壤数据上报记录。",
   "气象数据": "分类查看气象数据上报记录。",
   "农事记录": "分类查看农事记录。",
   "病虫上报": "分类查看病虫上报记录。",
   "蜜桔销售": "分类查看蜜桔销售上报记录。"},
}

MODAL_DESC = {
 "主体": "展示果园主体概况，含果园数量、覆盖面积、植株规模等统计指标。",
 "活跃度": "展示系统/设备活跃度，含在线设备数、数据更新频率等。",
 "植株": "展示植株监测概况，含监测株数与优/良/一般/异常等级分布。",
 "报告": "展示评价报告统计，含报告数量与生成趋势。",
 "评分": "展示综合评分分布与等级占比。",
 "农事": "展示农事建议的采纳与执行情况统计。",
 "采纳率": "展示建议采纳率趋势与对比。",
 "学习": "展示模型自学习/校准次数与效果统计。",
}

manifest = json.load(open(os.path.join(SHOT, "manifest.json"), encoding="utf-8"))
n = 0
for rel, labels in manifest.items():
    if rel.startswith("__"):
        continue
    for i in range(1, len(labels)):
        lab = labels[i]
        cap = "图：%s - %s" % (rel, lab)
        cp = find_para_by_text(cap)
        if not cp:
            print("WARN caption not found:", cap); continue
        img = image_after(cp)
        desc = TAB_DESC.get(rel, {}).get(lab)
        if desc:
            insert_after(img, desc, size=9.5); n += 1

modals = manifest.get("__growth_dashboard_modal", [])
for m in modals:
    cap = "图：3.3.1 生长模型监测大屏 - KPI弹窗（%s）" % m
    cp = find_para_by_text(cap)
    if not cp:
        print("WARN modal caption not found:", cap); continue
    img = image_after(cp)
    desc = MODAL_DESC.get(m)
    if desc:
        insert_after(img, desc, size=9.5); n += 1

doc.save(DOCX)
print("descriptions inserted:", n, "->", DOCX)
