# -*- coding: utf-8 -*-
"""仅修复需求说明书中的截图：按文档顺序将每张图映射到其"对应文件"页面，
用正确分辨率 PNG 替换（后台/大屏 1920x1080，移动端 430x932），并清理残留旧图。
不动第一章与缺失页面（已在上一轮完成）。"""
import os, copy, glob
from docx import Document
from docx.shared import Cm, RGBColor
from docx.oxml.ns import qn

PATH = r"D:\dev\pro_orangemodel\docs\南丰蜜桔模型管理系统-需求说明书.docx"
SHOT = r"D:\dev\pro_orangemodel\docs\screenshots"

doc = Document(PATH)
body = doc.element.body

shot_files = glob.glob(os.path.join(SHOT, "*.png"))
shot_names = {os.path.basename(f) for f in shot_files}
def shot_for(page):
    name = page.replace("/", "__").rsplit(".", 1)[0] + ".png"
    if name in shot_names:
        return os.path.join(SHOT, name)
    if page.endswith("growth_model_memo.html"):
        return shot_for("backhand/growth/growth_model.html")
    if page.endswith("pest_model_memo.html"):
        return shot_for("backhand/pest/pest_model.html")
    return None

replaced = 0
missing = 0
current_page = None
for el in list(body):
    if el.tag == qn('w:p'):
        txt = "".join(n.text or "" for n in el.iter(qn('w:t')))
        if "对应文件：" in txt:
            current_page = txt.split("对应文件：", 1)[1].strip().split()[0]
        if el.find('.//'+qn('w:drawing')) is not None and current_page:
            png = shot_for(current_page)
            if not png:
                missing += 1
                continue
            width_cm = 6.8 if current_page.startswith("mobile/") else 15.0
            tmp = doc.add_paragraph()
            run = tmp.add_run()
            run.add_picture(png, width=Cm(width_cm))
            new_drawing = tmp._p.find('.//'+qn('w:drawing'))
            old = el.find('.//'+qn('w:drawing'))
            old.getparent().replace(old, copy.deepcopy(new_drawing))
            tmp._p.getparent().remove(tmp._p)
            replaced += 1

print("imgs replaced=", replaced, "missing=", missing)

# 清理未被引用的旧图片关系
referenced = set()
for drawing in body.iter(qn('w:drawing')):
    blip = drawing.find('.//'+qn('a:blip'))
    if blip is not None:
        referenced.add(blip.get(qn('r:embed')))
dropped = 0
for rid in list(doc.part.rels):
    rel = doc.part.rels[rid]
    if rid not in referenced and rel.reltype.endswith("/image"):
        doc.part.drop_rel(rid)
        dropped += 1
print("orphan image rels dropped=", dropped)

# 全文档置黑
for p in doc.paragraphs:
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)

doc.save(PATH)
print("SAVED:", PATH)
