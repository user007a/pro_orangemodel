# -*- coding: utf-8 -*-
"""生成两份合同附件 Word 文档：软件设计方案 + 软件技术方案（黑白色、含架构图、约10页）。"""
import os, tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
OUT = r"D:\dev\pro_orangemodel\docs"
ASSET = tempfile.gettempdir() + r"\nf_diagrams"
os.makedirs(ASSET, exist_ok=True)

# ===================== 字体/样式辅助 =====================
def set_font(run, size=10.5, bold=False, name="宋体", eastasia="宋体", color=BLACK):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), eastasia)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)

def add_para(doc, text="", size=10.5, bold=False, align="left", indent=True,
             space_after=4, space_before=0, name="宋体", color=BLACK, line=1.5):
    p = doc.add_paragraph()
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT, "just": WD_ALIGN_PARAGRAPH.JUSTIFY}[align]
    pf = p.paragraph_format
    pf.line_spacing = line
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if indent and align in ("left", "just"):
        pf.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, name=name, color=color)
    return p

def add_heading(doc, text, level=1):
    sizes = {1: 15, 2: 12.5, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    set_font(r, size=sizes.get(level, 11), bold=True, name="黑体", eastasia="黑体")
    return p

def add_bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("• " + it)
        set_font(r, size=size)

def add_table(doc, headers, rows, widths=None, header_size=9.5, body_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=header_size, bold=True, name="黑体", eastasia="黑体")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=body_size)
            if widths and i < len(widths):
                cells[i].width = Cm(widths[i])
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_image(doc, path, width_cm=15.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(4)

def add_page_break(doc):
    doc.add_page_break()

def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "宋体"
    st.font.size = Pt(10.5)
    st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), "宋体"); rf.set(qn("w:ascii"), "宋体"); rf.set(qn("w:hAnsi"), "宋体")

def add_cover(doc, title, subtitle):
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    add_para(doc, title, size=22, bold=True, align="center", indent=False, space_after=10, name="黑体", color=BLACK)
    add_para(doc, subtitle, size=14, align="center", indent=False, space_after=4, color=BLACK)
    add_para(doc, "（合同附件）", size=12, align="center", indent=False, space_after=50, color=BLACK)
    add_para(doc, "委托方（甲方）：________________________", size=11, align="center", indent=False, space_after=6, color=BLACK)
    add_para(doc, "承建方（乙方）：________________________", size=11, align="center", indent=False, space_after=6, color=BLACK)
    add_para(doc, "文档版本：V1.0", size=11, align="center", indent=False, space_after=6, color=BLACK)
    add_para(doc, "编制日期：2026 年 __ 月 __ 日", size=11, align="center", indent=False, space_after=6, color=BLACK)
    add_page_break(doc)

def add_toc(doc, items):
    add_para(doc, "目  录", size=14, bold=True, align="center", indent=False, space_after=8, name="黑体", color=BLACK)
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(it)
        set_font(r, size=10.5)
    add_page_break(doc)

# ===================== 架构图（matplotlib，黑白） =====================
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle
import matplotlib.font_manager as fm

_FP = fm.FontProperties(fname="C:/Windows/Fonts/simhei.ttf")
plt.rcParams["font.family"] = _FP.get_name()
plt.rcParams["axes.unicode_minus"] = False

def _box(ax, x, y, w, h, text, fs=8.5):
    ax.add_patch(Rectangle((x, y), w, h, fill=True, facecolor="white", edgecolor="black", linewidth=1.1))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs, color="black", fontproperties=_FP)

def draw_layered(path, title, layers):
    n = len(layers)
    row_h = 1.05
    gap = 0.28
    fig_h = n * (row_h + gap) + 0.9
    fig, ax = plt.subplots(figsize=(9.8, fig_h))
    ax.set_xlim(0, 10); ax.set_ylim(0, n * (row_h + gap) + 0.7)
    ax.axis("off")
    ax.text(5, n * (row_h + gap) + 0.35, title, ha="center", va="center",
            fontsize=13, fontweight="bold", color="black", fontproperties=_FP)
    for i, (label, boxes) in enumerate(layers):
        top = n * (row_h + gap) - i * (row_h + gap)
        y = top - row_h
        ax.add_patch(Rectangle((0.15, y), 1.5, row_h, fill=True, facecolor="white", edgecolor="black", linewidth=1.1))
        ax.text(0.9, y + row_h / 2, label, ha="center", va="center", fontsize=9.5,
                fontweight="bold", color="black", fontproperties=_FP)
        nb = len(boxes)
        x0 = 1.85
        avail = 10 - x0 - 0.15
        bw = avail / nb
        for j, b in enumerate(boxes):
            bx = x0 + j * bw
            _box(ax, bx + 0.05, y, bw - 0.1, row_h, b, fs=8.3)
        if i < n - 1:
            next_top = n * (row_h + gap) - (i + 1) * (row_h + gap)
            ax.annotate("", xy=(5, next_top + row_h + 0.02), xytext=(5, y - 0.02),
                        arrowprops=dict(arrowstyle="->", color="black", lw=1.1))
    fig.savefig(path, dpi=200, bbox_inches="tight", pad_inches=0.08)
    plt.close()

def draw_flow(path, title, stages, note):
    n = len(stages)
    fig, ax = plt.subplots(figsize=(9.8, 3.4))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.2)
    ax.axis("off")
    ax.text(5, 3.0, title, ha="center", va="center", fontsize=13, fontweight="bold", color="black", fontproperties=_FP)
    x0, x1, y, h = 0.3, 9.7, 1.35, 0.85
    w = (x1 - x0 - (n - 1) * 0.22) / n
    for i, s in enumerate(stages):
        bx = x0 + i * (w + 0.22)
        _box(ax, bx, y, w, h, s, fs=8.3)
        if i < n - 1:
            ax.annotate("", xy=(bx + w + 0.2, y + h / 2), xytext=(bx + w + 0.02, y + h / 2),
                        arrowprops=dict(arrowstyle="->", color="black", lw=1.1))
    last = x0 + (n - 1) * (w + 0.22)
    ax.annotate("", xy=(x0 + w / 2, y - 0.02), xytext=(last + w / 2, y - 0.02),
                arrowprops=dict(arrowstyle="->", color="black", lw=1.1, connectionstyle="arc3,rad=-0.32"))
    ax.text(5, 0.45, note, ha="center", va="center", fontsize=8.5, color="black", fontproperties=_FP)
    fig.savefig(path, dpi=200, bbox_inches="tight", pad_inches=0.08)
    plt.close()

# 生成图
p_func = os.path.join(ASSET, "func_arch.png")
p_flow = os.path.join(ASSET, "core_flow.png")
p_tech = os.path.join(ASSET, "tech_arch.png")

draw_layered(p_func, "总体功能架构图", [
    ("用户与接入", ["后台管理员", "种植农户\n(小程序)", "大屏访客", "第三方系统"]),
    ("应用子系统", ["模型管理后台\n(Web)", "模型监测大屏\n(可视化)", "蜜桔农业助手\n(微信小程序)", "第三方对接"]),
    ("核心业务引擎", ["生长模型引擎", "病虫害模型引擎", "预警管理", "数据管理"]),
    ("数据资源", ["基础数据库\n(9类)", "外部数据源\n(气象/农业/价格/测报)"]),
    ("基础设施", ["服务器 / 网络 / 安全"]),
])
draw_flow(p_flow, "核心业务闭环图",
          ["数据采集", "因子归集", "模型计算", "诊断分析", "评价/预测报告", "预警生成", "农事/防控建议", "建议采纳"],
          "农户采纳建议后回灌数据，形成持续优化的业务闭环")
draw_layered(p_tech, "技术总体架构图", [
    ("展示层", ["Web后台\n(Vue)", "监测大屏\n(ECharts)", "微信小程序", "第三方门户"]),
    ("网关层", ["API网关", "统一认证\n(SSO)"]),
    ("业务服务层", ["生长模型服务", "病虫害模型服务", "数据服务", "预警服务", "资讯/用户服务"]),
    ("数据层", ["业务数据库\n(MySQL)", "缓存\n(Redis)", "文件/图片存储", "时序/气象数据"]),
    ("集成层", ["设备IoT接入", "外部数据接口", "消息推送"]),
    ("基础设施", ["云服务器/容器", "Nginx", "备份与监控"]),
])

# ===================== 文档一：软件设计方案 =====================
def build_design():
    doc = Document()
    set_base_style(doc)
    # 页边距
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.0)
    add_cover(doc, "南丰蜜桔模型管理系统", "软件设计方案")
    add_toc(doc, [
        "1  项目概述", "2  总体功能架构", "3  角色与权限设计",
        "4  功能模块详细设计", "5  关键业务流程", "6  界面与交互设计",
        "7  非功能性要求概述", "8  验收与交付说明",
    ])

    # 1 项目概述
    add_heading(doc, "1  项目概述")
    add_heading(doc, "1.1  项目背景", 2)
    add_para(doc, "南丰蜜桔是区域性特色优势农产品，其产量与品质高度依赖生长环境与病虫害防控水平。传统管理模式依赖人工经验，存在数据分散、评价主观、预警滞后等问题。南丰蜜桔模型管理系统是“南丰县蜜桔全产业链智慧建设项目”的重要组成部分，本次在现有系统基础上进行完善与整改，以生长模型与病虫害模型为核心，打通数据采集、模型计算、诊断评价、预警推送与农事指导的业务闭环。")
    add_heading(doc, "1.2  建设目标", 2)
    add_bullets(doc, [
        "建立统一的南丰蜜桔生长与病虫害模型管理能力，实现关键因子的数字化采集与归集；",
        "提供生长综合评价与病虫害风险预测，输出可执行的农事建议与防控方案；",
        "通过管理后台、监测大屏与微信小程序三端协同，服务管理人员与广大种植农户；",
        "对接设备、气象及行业外部数据，形成可持续优化的数据底座。",
    ])
    add_heading(doc, "1.3  建设范围", 2)
    add_para(doc, "本系统由四个相对独立又相互协同的子系统构成，详细功能以《南丰蜜桔模型管理系统项目-功能清单》为准：")
    add_table(doc, ["子系统", "定位", "主要服务对象"], [
        ["模型管理后台", "Web 管理端，承载模型配置、数据管理与系统设置", "管理员、业务人员"],
        ["模型监测大屏", "可视化展示端，呈现生长与病虫害核心指标", "决策与展示访客"],
        ["蜜桔农业助手", "微信小程序，面向农户提供识别、评价与资讯", "种植农户"],
        ["第三方对接", "对接单点登录、设备与外部数据来源", "第三方系统"],
    ], widths=[3.2, 8.5, 3.3])
    add_heading(doc, "1.4  术语定义", 2)
    add_table(doc, ["术语", "说明"], [
        ["生长模型", "基于生长因子对植株生长状态进行适宜度计算与综合评分的模型"],
        ["病虫害模型", "基于监测与参数对病虫害发生风险进行预测与分级的模型"],
        ["因子", "影响生长或病虫害的关键环境/农事指标（如温湿度、物候等）"],
        ["农事建议", "模型依据评价结果推荐的栽培管理措施"],
        ["预警", "当指标或风险达到阈值时生成的提示信息"],
    ], widths=[3.0, 12.0])
    add_page_break(doc)

    # 2 总体功能架构
    add_heading(doc, "2  总体功能架构")
    add_heading(doc, "2.1  架构总览", 2)
    add_para(doc, "系统采用“用户接入—应用子系统—核心业务引擎—数据资源—基础设施”的分层组织方式，各层职责清晰、自下而上提供支撑。总体功能架构如下图所示。")
    add_image(doc, p_func, width_cm=15.5)
    add_heading(doc, "2.2  子系统定位", 2)
    add_para(doc, "模型管理后台负责模型与数据的集中管理；模型监测大屏负责结果的可视化呈现；蜜桔农业助手负责农户侧的交互与服务；第三方对接负责与外部系统及设备的贯通。四者围绕核心业务引擎协同运作。")
    add_heading(doc, "2.3  核心业务引擎（重点）", 2)
    add_para(doc, "本系统的核心价值集中在两大模型引擎：生长模型引擎与病虫害模型引擎，并辅以预警管理与数据管理，构成系统的“大脑”。")
    add_bullets(doc, [
        "生长模型引擎：围绕适宜度、门控、生长增量、扣分与综合评分进行计算，输出植株生长评价；",
        "病虫害模型引擎：基于监测数据与预测参数输出风险等级与防控建议；",
        "预警管理：统一承接两类模型的异常信号，按规则生成、分级与推送预警；",
        "数据管理：汇聚基础数据库与外部数据源，为模型提供可靠输入。",
    ])
    add_heading(doc, "2.4  与外部系统的关系", 2)
    add_para(doc, "系统通过第三方对接获取设备实时数据、气象局与农业大数据、价格行情及测报等信息，并将用户与消息能力延伸至微信小程序与第三方门户，形成开放的数据与能力通道。")
    add_page_break(doc)

    # 3 角色与权限
    add_heading(doc, "3  角色与权限设计")
    add_heading(doc, "3.1  用户角色", 2)
    add_table(doc, ["角色", "说明"], [
        ["系统管理员", "负责系统用户、角色与全局配置"],
        ["业务管理员", "负责模型、数据、预警与资讯的日常运营"],
        ["种植农户", "通过微信小程序使用识别、评价、资讯与上报功能"],
        ["大屏访客", "查看监测大屏展示内容"],
        ["第三方系统", "通过接口进行用户、设备与数据对接"],
    ], widths=[3.2, 11.8])
    add_heading(doc, "3.2  权限设计原则", 2)
    add_para(doc, "系统采用基于角色的访问控制（RBAC）思路：角色与权限解耦，用户通过分配角色获得相应菜单与操作权限；后台管理端与微信后台均遵循统一的权限模型，确保权限可控、可审计。")
    add_heading(doc, "3.3  角色-功能权限对应（示意）", 2)
    add_table(doc, ["角色", "模型管理后台", "监测大屏", "微信小程序", "系统管理"], [
        ["系统管理员", "√", "查看", "—", "√"],
        ["业务管理员", "√", "查看", "—", "部分"],
        ["种植农户", "—", "—", "√", "—"],
        ["大屏访客", "—", "查看", "—", "—"],
    ], widths=[3.0, 3.5, 3.0, 3.0, 2.5])
    add_page_break(doc)

    # 4 功能模块详细设计
    add_heading(doc, "4  功能模块详细设计")
    add_para(doc, "以下按子系统列出一级模块与主要功能要点，二级菜单与内容描述详见《功能清单》附件。")
    add_heading(doc, "4.1  模型管理后台", 2)
    add_table(doc, ["一级模块", "主要功能"], [
        ["系统首页", "系统框架、登录认证、总览看板"],
        ["生长模型", "模型说明、基础档案、因子采集、农事建议、模型计算、诊断分析、评价报告、数据校准"],
        ["病虫害模型", "模型说明、监测数据、参数配置、预测计算、风险诊断、预测报告、防控方案"],
        ["数据管理", "病害/虫害/灾害/品种/气象/墒情/视频/价格/物候数据库、外部抓取"],
        ["预警管理", "预警管理、处置记录、预警规则、预警模板"],
        ["微信后台", "微信用户、果园管理、资讯管理、人工上报、建议采纳、识别日志"],
        ["系统管理", "用户管理、角色管理"],
    ], widths=[3.2, 11.8])
    add_heading(doc, "4.2  模型监测大屏", 2)
    add_table(doc, ["一级模块", "主要功能"], [
        ["生长模型大屏", "KPI概览、环境监测、植株综合评测、农事建议、监测地图、子项分析、落地统计、阶段分布"],
        ["病虫害模型大屏", "KPI概览、高发排名、实时识别、分级预警、风险热力图、气象数据、灾害预判、准确率趋势、损失对比、防护落地"],
    ], widths=[3.2, 11.8])
    add_heading(doc, "4.3  蜜桔农业助手（微信小程序）", 2)
    add_table(doc, ["一级模块", "主要功能"], [
        ["首页", "综合看板：天气、预警、消息、服务入口、价格、农事推荐、资讯"],
        ["病虫害服务", "AI智能识别、病虫害知识库"],
        ["长势评价", "植株指标录入、综合评价报告"],
        ["数据采集", "环境监测、数据上报、我的果园"],
        ["农技资讯", "资讯首页、列表、详情、学习中心"],
        ["价格行情", "价格指数与历史趋势"],
        ["消息中心", "预警/农事/AI/资讯消息与反馈"],
        ["个人中心", "个人信息、上报记录、我的报告、设置"],
    ], widths=[3.2, 11.8])
    add_heading(doc, "4.4  第三方对接", 2)
    add_table(doc, ["一级模块", "主要功能"], [
        ["平台对接", "用户与单点登录、设备数据对接、消息回调与业务数据对接"],
    ], widths=[3.2, 11.8])
    add_page_break(doc)

    # 5 关键业务流程
    add_heading(doc, "5  关键业务流程")
    add_heading(doc, "5.1  生长评价主闭环（重点）", 2)
    add_para(doc, "系统以“数据驱动评价”为主线，形成如下闭环：")
    add_image(doc, p_flow, width_cm=15.5)
    add_para(doc, "即：通过物联网与人工采集获取生长因子，归集后由生长模型计算适宜度与综合评分，经诊断分析生成评价报告，并据结果产生农事建议；农户采纳建议后回灌数据，持续优化模型。")
    add_heading(doc, "5.2  病虫害防控流程", 2)
    add_bullets(doc, [
        "监测/识别：采集监测数据或由农户拍照 AI 识别；",
        "预测计算：病虫害模型依据参数输出预测结果；",
        "风险诊断：评估风险等级并给出诊断结论；",
        "预警生成：达到阈值时生成分级预警；",
        "防控方案：输出针对性防控建议；",
        "处置记录：记录处置措施、结果与责任人。",
    ])
    add_heading(doc, "5.3  预警触发与处置", 2)
    add_para(doc, "预警由规则引擎统一触发，支持按指标阈值设定等级与通知方式；预警生成后进入处置流程，处置记录可追溯，形成“触发—通知—处置—闭环”的完整链路。")
    add_heading(doc, "5.4  数据流转概述", 2)
    add_para(doc, "外部数据源与设备数据进入基础数据库，经归集为模型因子；模型输出评价、预测与预警，再经大屏与小程序对外呈现，农户反馈回灌至数据层，构成端到端的数据流转。")
    add_page_break(doc)

    # 6 界面与交互
    add_heading(doc, "6  界面与交互设计")
    add_heading(doc, "6.1  三端设计风格", 2)
    add_table(doc, ["端", "风格基调", "设计要点"], [
        ["模型管理后台", "专业、严谨", "清晰的信息架构、表格与表单为主、强调操作效率"],
        ["模型监测大屏", "可视化、直观", "图表与地图为主、突出核心指标与态势"],
        ["微信小程序", "轻量、易用", "移动优先、流程简化、面向农户高频场景"],
    ], widths=[3.2, 3.5, 8.3])
    add_heading(doc, "6.2  主要页面说明", 2)
    add_bullets(doc, [
        "后台首页：系统总览看板，集中展示模型运行状态与核心统计；",
        "诊断分析页：六维度因子分析与扣分溯源，支撑评价解读；",
        "大屏：KPI、地图与趋势图组合，便于宏观研判；",
        "小程序首页：天气、预警与常用服务一站式入口。",
    ])
    add_heading(doc, "6.3  响应与适配", 2)
    add_para(doc, "后台与大屏面向桌面浏览器，小程序面向移动端；各端遵循统一的设计语言与配色规范，保证视觉一致性与操作连贯性。")
    add_page_break(doc)

    # 7 非功能性
    add_heading(doc, "7  非功能性要求概述")
    add_para(doc, "以下为系统建设的总体非功能性要求，具体指标在执行阶段结合实际情况确定，不作为对开发实现的硬性约束。")
    add_heading(doc, "7.1  可用性", 2)
    add_para(doc, "系统应保证业务时段稳定可用，关键服务具备基本的容错与恢复能力。")
    add_heading(doc, "7.2  性能", 2)
    add_para(doc, "在满足日常业务并发与数据规模的前提下，各主要功能操作响应顺畅；模型计算与报表生成在合理时间内完成。")
    add_heading(doc, "7.3  兼容性", 2)
    add_para(doc, "后台与大屏兼容主流桌面浏览器；小程序适配常见移动设备与微信运行环境。")
    add_heading(doc, "7.4  数据安全与隐私", 2)
    add_para(doc, "对用户数据与业务数据实行分级保护，传输与存储采取必要的安全措施，遵循相关数据安全与隐私保护要求。")
    add_heading(doc, "7.5  可维护性与可扩展性", 2)
    add_para(doc, "系统采用模块化设计，便于功能扩展与后续迭代；配置与代码分离，降低维护成本。")
    add_page_break(doc)

    # 8 验收与交付
    add_heading(doc, "8  验收与交付说明")
    add_heading(doc, "8.1  验收依据", 2)
    add_para(doc, "以《南丰蜜桔模型管理系统项目-功能清单》及本方案约定的功能范围为验收依据，逐子系统、逐模块核对实现情况。")
    add_heading(doc, "8.2  交付物", 2)
    add_bullets(doc, [
        "可运行的系统（后台、大屏、小程序与对接）；",
        "《软件技术方案》及配套设计文档；",
        "用户操作手册与必要的培训材料；",
        "数据字典与接口说明（详见技术方案）。",
    ])
    add_heading(doc, "8.3  文档约定", 2)
    add_para(doc, "本《软件设计方案》与配套的《软件技术方案》共同作为合同附件；设计方案界定“做什么”，技术方案界定“怎么做”，两者范围一致、互不冲突。")
    out = os.path.join(OUT, "南丰蜜桔模型管理系统-软件设计方案.docx")
    doc.save(out)
    print("saved:", out)

# ===================== 文档二：软件技术方案 =====================
def build_tech():
    doc = Document()
    set_base_style(doc)
    for s in doc.sections:
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.0)
    add_cover(doc, "南丰蜜桔模型管理系统", "软件技术方案")
    add_toc(doc, [
        "1  项目概述", "2  技术选型与总体架构", "3  部署架构", "4  数据架构与库设计",
        "5  接口与集成设计", "6  模型计算与算法实现", "7  安全设计",
        "8  性能与可靠性", "9  开发与交付方案", "10  运维与培训",
    ])

    # 1 项目概述（参考整改实施方案）
    add_heading(doc, "1  项目概述")
    add_heading(doc, "1.1  项目背景", 2)
    add_para(doc, "南丰蜜桔模型管理系统是“南丰县蜜桔全产业链智慧建设项目”的重要组成部分。根据项目技术审查提出的整改意见，并结合实际建设范围，本次建设在现有系统基础上进行完善与整改，重点提升生长模型、病虫害模型及相关数据能力，使系统更好地服务于蜜桔生产管理与产业服务。")
    add_heading(doc, "1.2  整改与建设原则", 2)
    add_para(doc, "本方案坚持“充分利用现有平台、设备和数据”的原则：不新建独立平台，不大规模新增物联网设备，所有完善内容均整合到现有模型管理后台、模型监测大屏、蜜桔农业助手与第三方对接能力之中，重点对现有功能进行补充、整合与打通。")
    add_heading(doc, "1.3  建设基础", 2)
    add_para(doc, "系统已在前期建设了模型管理后台（Web 管理端）、模型监测大屏（可视化端）、蜜桔农业助手（微信小程序）与第三方对接能力。本技术方案在上述既有能力基础上开展功能完善、集成与运维设计。")
    add_heading(doc, "1.4  建设内容与范围", 2)
    add_para(doc, "围绕以下方向开展完善，详细功能以《南丰蜜桔模型管理系统项目-功能清单》为准：")
    add_bullets(doc, [
        "生长模型完善：完善模型说明、基础档案、因子采集、模型计算、诊断分析、农事建议与评价报告；",
        "病虫害模型完善：完善监测数据、参数配置、预测计算、风险诊断、预测报告与防控方案，并增强图片辅助识别能力；",
        "数据与预警完善：完善病害/虫害/灾害/品种/气象/墒情/价格等基础数据库及外部数据抓取，强化预警规则与处置闭环；",
        "销售与市场数据作为数据管理能力的延伸，通过价格数据库与外部网站采集实现价格与行情监测。",
    ])
    add_para(doc, "实施过程中引入农技人员参与模型内容审核与试运行确认，确保评价结论与农事建议符合南丰蜜桔本地生产实际。")
    add_heading(doc, "1.5  实施周期与试点", 2)
    add_para(doc, "建设周期控制在 2—3 个月，选择少量具有代表性的果园开展试点应用，经系统测试、人员培训与验收资料编制后上线使用。")
    add_page_break(doc)

    # 2
    add_heading(doc, "2  技术选型与总体架构")
    add_heading(doc, "2.1  技术选型原则", 2)
    add_para(doc, "在满足业务需求的前提下，优先采用成熟、稳定、社区活跃的主流技术，兼顾开发效率、可维护性与可扩展性，降低长期运维风险。")
    add_heading(doc, "2.2  总体技术架构", 2)
    add_para(doc, "本方案对应并落地《软件设计方案》的功能范围，技术总体架构如下图所示，自顶向下分为展示层、网关层、业务服务层、数据层、集成层与基础设施。")
    add_image(doc, p_tech, width_cm=15.5)
    add_heading(doc, "2.3  技术栈清单", 2)
    add_table(doc, ["层次", "技术选型（示意）"], [
        ["前端", "后台：Vue 系列框架 + 组件库；大屏：可视化图表库；小程序：微信原生/框架"],
        ["后端", "主流服务端语言与框架（如 Java/Python/Node 等），提供 RESTful 接口"],
        ["数据库", "关系型数据库（业务数据）+ 缓存（热点数据）+ 文件/对象存储（图片）"],
        ["部署", "Web 服务器与反向代理（如 Nginx），可运行于云服务器或容器环境"],
        ["集成", "消息推送、定时任务、外部 HTTP/接口对接"],
    ], widths=[3.0, 12.0])
    add_heading(doc, "2.4  与设计方案的对应关系", 2)
    add_para(doc, "技术方案的模块划分与《软件设计方案》四大子系统一一对应：后台对应业务服务与数据管理，大屏对应可视化展示，小程序对应移动端展示层，第三方对接对应集成层，确保两份文档范围一致、实现可追溯。")
    add_page_break(doc)

    # 3
    add_heading(doc, "3  部署架构")
    add_heading(doc, "3.1  部署拓扑", 2)
    add_para(doc, "系统采用前后端分离部署：前端静态资源经反向代理提供服务，后端服务与数据库独立部署，集成层通过安全通道对接外部系统与设备。")
    add_heading(doc, "3.2  运行环境", 2)
    add_table(doc, ["环境", "用途"], [
        ["开发环境", "功能开发与联调"],
        ["测试环境", "功能、集成与验收测试"],
        ["生产环境", "正式对外提供服务"],
    ], widths=[3.5, 11.5])
    add_heading(doc, "3.3  高可用与备份", 2)
    add_para(doc, "生产环境对数据库与关键服务采取必要的冗余与备份策略，定期备份业务数据，保障故障情况下可恢复。")
    add_page_break(doc)

    # 4
    add_heading(doc, "4  数据架构与库设计")
    add_heading(doc, "4.1  数据架构概述", 2)
    add_para(doc, "数据层由业务数据库、缓存、文件/对象存储及时序/气象数据构成。业务数据库承载模型、因子、地块、植株、预警、用户、设备与资讯等核心实体。")
    add_heading(doc, "4.2  核心数据实体", 2)
    add_bullets(doc, [
        "模型与因子：生长/病虫害模型定义、因子采集与计算记录；",
        "地块与植株：果园地块档案、植株生长档案与监测数据；",
        "预警：预警记录、规则、模板与处置记录；",
        "用户与设备：后台用户、微信用户、设备与果园绑定；",
        "资讯与价格：农技资讯、价格行情与物候等基础数据。",
    ])
    add_heading(doc, "4.3  主要数据表（示意）", 2)
    add_table(doc, ["分类", "示例表"], [
        ["模型", "生长模型表、病虫害模型表、因子采集表、模型计算日志表"],
        ["评价/预警", "评价报告表、预警记录表、预警规则表、处置记录表"],
        ["资源", "地块表、植株表、病害库、虫害库、灾害库、品种库"],
        ["监测", "气象站点表、墒情表、视频设备表、价格表、物候表"],
        ["用户", "系统用户表、角色表、微信用户表、果园表"],
    ], widths=[3.0, 12.0])
    add_heading(doc, "4.4  数据流转与存储", 2)
    add_para(doc, "外部数据与设备数据经集成层写入业务数据库，模型计算读取因子并写回评价与预警结果；图片等附件存入文件/对象存储，热点数据入缓存以提升访问效率。")
    add_page_break(doc)

    # 5
    add_heading(doc, "5  接口与集成设计")
    add_heading(doc, "5.1  内部接口规范", 2)
    add_para(doc, "系统内部服务间及前后端之间通过 RESTful 风格接口通信，采用统一的请求/响应结构、状态码与鉴权方式，便于联调与维护。")
    add_heading(doc, "5.2  第三方对接", 2)
    add_table(doc, ["对接对象", "内容"], [
        ["用户与单点登录", "用户体系打通、SSO 登录、消息回调与业务数据对接"],
        ["设备（IoT）", "多种类型设备的实时数据接入与设备管理"],
        ["外部数据", "气象局、农业大数据、价格行情、测报等系统数据抓取与对接"],
    ], widths=[4.0, 11.0])
    add_heading(doc, "5.3  接口安全", 2)
    add_para(doc, "对外接口需进行身份认证与权限校验，敏感数据传输加密，并对调用频率与异常访问做必要控制。")
    add_page_break(doc)

    # 6
    add_heading(doc, "6  模型计算与算法实现")
    add_heading(doc, "6.1  生长模型计算", 2)
    add_para(doc, "生长模型以因子采集数据为输入，依次进行适宜度计算、门控判断、生长增量累计、扣分处理，最终形成综合评分。模型计算中心负责批量调度并保留计算日志，支持数据校准与参数微调。")
    add_heading(doc, "6.2  病虫害预测模型", 2)
    add_para(doc, "病虫害模型基于监测数据与配置的参数（阈值、预测参数）进行预测计算，输出风险等级与诊断结论，并生成预测报告与防控方案。")
    add_heading(doc, "6.3  AI 图像识别", 2)
    add_para(doc, "针对农户拍照上传的病虫害图片，提供智能识别能力，返回病害/虫害名称、特征与防治建议，识别结果记入日志并支撑风险诊断。")
    add_heading(doc, "6.4  计算调度", 2)
    add_para(doc, "模型计算以任务方式统一调度，计算过程可监控、可回溯，异常计算记录日志以便排查，保障评价与预警的时效性。")
    add_page_break(doc)

    # 7
    add_heading(doc, "7  安全设计")
    add_heading(doc, "7.1  身份认证与授权", 2)
    add_para(doc, "系统采用统一认证，后台与接口均需登录鉴权；结合 RBAC 实现菜单与操作级权限控制，第三方接入使用受控的凭证机制。")
    add_heading(doc, "7.2  传输与存储安全", 2)
    add_para(doc, "关键链路采用加密传输，口令等敏感信息加密存储，文件与图片按权限受控访问。")
    add_heading(doc, "7.3  权限控制", 2)
    add_para(doc, "后台管理端按角色分配菜单与操作权限，微信后台按用户隔离其果园与数据，确保数据“看得见的才可用”。")
    add_heading(doc, "7.4  日志审计", 2)
    add_para(doc, "对关键操作、登录与模型计算保留日志，支持事后审计与问题追溯。")
    add_heading(doc, "7.5  数据安全合规", 2)
    add_para(doc, "遵循相关数据安全和隐私保护要求，对个人信息与业务数据实行分级管理与最小必要使用。")
    add_page_break(doc)

    # 8
    add_heading(doc, "8  性能与可靠性")
    add_heading(doc, "8.1  性能设计", 2)
    add_para(doc, "通过前后端分离、缓存、异步计算与数据库优化等手段，保障日常业务并发下的操作响应与报表生成效率；具体性能指标在执行阶段结合实际规模确定。")
    add_heading(doc, "8.2  可靠性与容灾", 2)
    add_para(doc, "对关键服务与数据采取冗余、备份与恢复机制，配合监控告警，降低故障影响，保障业务连续性。")
    add_page_break(doc)

    # 9
    add_heading(doc, "9  开发与交付方案")
    add_heading(doc, "9.1  开发语言与框架", 2)
    add_para(doc, "依据技术选型采用成熟的服务端语言与前端框架进行开发，代码结构清晰、分层合理，便于团队协作与后续维护。")
    add_heading(doc, "9.2  代码与配置管理", 2)
    add_para(doc, "使用版本控制系统管理源码，环境配置与代码分离，建立规范的分支与发布流程。")
    add_heading(doc, "9.3  测试策略", 2)
    add_table(doc, ["测试类型", "目的"], [
        ["单元测试", "验证各模块逻辑正确性"],
        ["集成测试", "验证服务间与内外接口协同"],
        ["验收测试", "对照功能清单核对实现范围"],
    ], widths=[3.5, 11.5])
    add_heading(doc, "9.4  交付物清单", 2)
    add_bullets(doc, [
        "可运行系统及部署说明；",
        "数据字典与接口说明文档；",
        "用户操作手册；",
        "必要的源码与构建产物（按合同约定）。",
    ])
    add_heading(doc, "9.5  实施里程碑（示意）", 2)
    add_table(doc, ["阶段", "交付内容"], [
        ["需求与设计", "需求确认、设计方案与技术方案"],
        ["开发实现", "各子系统功能开发与联调"],
        ["测试验收", "测试、试点与验收"],
        ["上线运维", "正式上线、培训与运维支持"],
    ], widths=[3.5, 11.5])
    add_page_break(doc)

    # 10
    add_heading(doc, "10  运维与培训")
    add_heading(doc, "10.1  运维监控", 2)
    add_para(doc, "对服务运行状态、接口调用与关键资源进行监控，出现异常及时告警，定期开展数据备份与健康检查。")
    add_heading(doc, "10.2  用户培训", 2)
    add_para(doc, "面向管理员与业务人员提供后台与大屏的操作培训，面向农户提供小程序使用引导，配套操作手册。")
    add_heading(doc, "10.3  技术支持", 2)
    add_para(doc, "提供约定周期内的技术支持与问题响应，保障系统稳定运行与持续优化。")
    out = os.path.join(OUT, "南丰蜜桔模型管理系统-软件技术方案.docx")
    doc.save(out)
    print("saved:", out)

if __name__ == "__main__":
    build_design()
    build_tech()
    print("ALL DONE")
